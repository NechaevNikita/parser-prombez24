from docx.enum.text import WD_COLOR_INDEX
from docx import Document
import requests
from bs4 import BeautifulSoup


url = 'https://prombez24.com/ticket/ordered?testId=169&page=0&size=257'
resource = requests.get(url)
soup = BeautifulSoup(resource.text, 'lxml')
questions = soup.find_all('div', class_='question row')

document = Document()
document.add_heading('Промышленная безопасность - https://prombez24.com', 1)


for i, question in enumerate(questions, start=1):
	question_text = question.find('div', class_='question__text').text.strip()
	# print(f'\nВопрос#{i}\n{question_text}')
	p = document.add_paragraph()
	p.add_run(f'\nВопрос# {i}').bold = True
	p.add_run(f'\n{question_text}')
	answers = question.find_all('div', class_='question__answers-list-item')
	for j, answer in enumerate(answers, start=1):
		answer_text = answer.find('span', class_='label').text.strip()
		value = answer.find_all('input')[-2]['value']
		# print(f'{j}.{value}. {answer_text}')

		if value == 'true':
			highlight_para = document.add_paragraph().add_run(f'{j}. {answer_text}').font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
		else:
			document.add_paragraph(f'{j}. {answer_text}')

document.save('Промышленная безопасность Б 1.10.docx')
